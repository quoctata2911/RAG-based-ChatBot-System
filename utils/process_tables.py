import time
import mammoth
from bs4 import BeautifulSoup
from docx.api import Document
import re

def extract_and_replace_docx_tables(docx_file, chunk_marker):
    start_time = time.time()  # Record start time

    document = Document(docx_file)
    docx_tables = document.tables
    total_tables = len(document.tables)

    with open(docx_file, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value 

    tables = extract_html_tables(html)
    
    html_chunked_tables = get_html_table_chunks(tables, chunk_marker=chunk_marker)
    
    
    html_tables = []
    for table in html_chunked_tables:
        temp_document = Document()
        html_table = temp_document.add_paragraph(table)._element
        html_table.alignment = 0
        html_tables.append(html_table)


    track = 0    
    while len(document.tables) > 0:
        track += 1
        try:
            html_table = html_tables[0]
            document.element.body.replace(document.tables[0]._element, html_table)
            html_tables.remove(html_table)
            end_time = time.time()  # Record end time
            # print(f'{track} of {total_tables} | Success | Time: {end_time - start_time:.2f} seconds')

        except Exception as e:
            print(f'{track} of {total_tables} | Fail: {e}')
            if track >= 200:
                break
            
    return document, html_chunked_tables

def extract_html_tables(html):
    soup = BeautifulSoup(html, 'html.parser')
    tables = soup.find_all('table')
    return tables

def get_html_table_chunks(tables, chunk_marker):
    
    html_chunk_marker = '<strong>' + chunk_marker + '</strong>'

    html_table_chunks = []
    
    for table_soup in tables:

        html_table_string = str(table_soup)
        html_table_string = html_table_string.replace('<table>', '<table>\n')
        html_table_string = html_table_string.replace('<tr>', '\n<tr>\n')
        html_table_string = html_table_string.replace('</tr>', '\n</tr>')
        html_table_string = html_table_string.replace('<thead>', '<thead>')
        html_table_string = html_table_string.replace('</thead>', '\n</thead>\n')
        html_table_string = html_table_string.replace('<tbody>', '<tbody>')
        html_table_string = html_table_string.replace('</tbody>', '\n</tbody>\n')

        with open('table_html.txt', mode='w', encoding='utf8') as f:
            f.write(html_table_string)

        with open('table_html.txt', mode='r', encoding='utf8') as f:
            lines = f.readlines()

        start_table = lines[0].strip()
        end_table = lines[-1].strip()
        
        # Get start and end tags for tbody
        start_tbody = '<tbody>' if '<tbody>' in html_table_string else ''
        end_tbody = '</tbody>' if '</tbody>' in html_table_string else ''

        # Extract and clean headers if present
        headers = str(table_soup.find('thead')) if 'thead' in html_table_string else ''
        headers = re.sub(r'>\n\s*<', '><', headers)

        processed_lines = []
        for line in lines:
            if chunk_marker in line:
                start_index = line.find(html_chunk_marker)
                chunk_start = start_index - len('<p>')
                chunk_end = start_index + len(html_chunk_marker) + len('</p>')

                chunk_html = line[chunk_start:chunk_end]
                if chunk_html.startswith('<p>') & chunk_html.endswith('</p>'):
                    line = line.replace('<p>', '')
                    line = line.replace('</p>', '')
                else:
                    pass

                line = line.replace(html_chunk_marker, '')
                line = line.replace(' </td>', '</td>').strip()
                line += chunk_marker

            processed_lines.append(line)

        processed_lines = [line.strip() for line in processed_lines]
        html_table = ''.join(processed_lines)
        html_chunks = html_table.split(chunk_marker)

        proccessed_html_chunks = []
        for index, chunk in enumerate(html_chunks):
            if index == 0:
                chunk += (end_tbody + end_table)
                first_chunk = chunk.replace(end_table, '')
                start = first_chunk.find('<tr>')
                end = first_chunk.find('</tr>') + len('</tr>')
                headers = first_chunk[start:end]

                    
            elif chunk == html_chunks[-1]:
                chunk = start_table + headers + start_tbody + chunk 

            else:
                chunk = start_table + headers + start_tbody + chunk + end_tbody + end_table

            proccessed_html_chunks.append(chunk)

        chunks_to_html = ''
        for html_chunk in proccessed_html_chunks:
                chunks_to_html += wrap_signal(html_chunk, signal_type='html')
                if html_chunk != proccessed_html_chunks[-1]:
                    chunks_to_html += f'\n\n{chunk_marker}\n\n'
        
        html_table_chunks.append(chunks_to_html)
        
    return html_table_chunks

def wrap_signal(data, signal_type):
    data = f"```{signal_type}\n{data}\n```"
    return data